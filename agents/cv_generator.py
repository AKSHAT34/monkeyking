"""
World-Class CV Generator — Creates detailed, ATS-optimized 2-page CVs.

Design principles:
- Single column, no tables, no images (ATS-safe)
- Clean typography with clear section headers
- Full detail: all bullets, all projects, all experience
- Tailored summary + skills per job description
- Professional spacing and formatting
- PDF and DOCX output
"""
import os
from datetime import datetime
from pathlib import Path
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable, KeepTogether
)
from reportlab.lib.colors import HexColor


class CVGenerator:
    """Generates world-class ATS-optimized 2-page CVs in PDF and DOCX."""

    OUTPUT_DIR = Path(__file__).parent.parent / "data" / "tailored_cvs"

    def __init__(self):
        self.OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        self._init_styles()

    def _init_styles(self):
        """Create professional typography styles."""
        base = getSampleStyleSheet()
        self.styles = {
            "name": ParagraphStyle("Name", parent=base["Title"],
                fontSize=20, leading=24, spaceAfter=1, alignment=TA_CENTER,
                textColor=HexColor("#1a1a1a"), fontName="Helvetica-Bold"),
            "title": ParagraphStyle("Title", parent=base["Normal"],
                fontSize=11, leading=14, textColor=HexColor("#555555"),
                alignment=TA_CENTER, spaceAfter=3),
            "contact": ParagraphStyle("Contact", parent=base["Normal"],
                fontSize=9, leading=12, alignment=TA_CENTER, spaceAfter=6,
                textColor=HexColor("#444444")),
            "section": ParagraphStyle("Section", parent=base["Heading2"],
                fontSize=12, leading=15, spaceBefore=10, spaceAfter=4,
                textColor=HexColor("#1a1a3e"), fontName="Helvetica-Bold",
                borderWidth=0, borderPadding=0),
            "subsection": ParagraphStyle("Subsection", parent=base["Normal"],
                fontSize=10, leading=13, spaceAfter=1, spaceBefore=4,
                textColor=HexColor("#222222"), fontName="Helvetica-Bold"),
            "meta": ParagraphStyle("Meta", parent=base["Normal"],
                fontSize=9, leading=12, spaceAfter=2,
                textColor=HexColor("#666666"), fontName="Helvetica-Oblique"),
            "body": ParagraphStyle("Body", parent=base["Normal"],
                fontSize=9.5, leading=12.5, spaceAfter=2,
                textColor=HexColor("#333333")),
            "bullet": ParagraphStyle("Bullet", parent=base["Normal"],
                fontSize=9, leading=12, leftIndent=14, spaceAfter=1.5,
                textColor=HexColor("#333333")),
            "skills": ParagraphStyle("Skills", parent=base["Normal"],
                fontSize=9, leading=12, spaceAfter=2,
                textColor=HexColor("#333333")),
        }

    def generate_pdf(self, profile: dict, job_title: str, company: str,
                     job_description: str = "", summary_type: str = "ai_product_manager") -> str:
        """Generate a detailed, tailored 2-page PDF CV."""
        safe_co = "".join(c if c.isalnum() else "_" for c in company)[:25]
        safe_t = "".join(c if c.isalnum() else "_" for c in job_title)[:25]
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"CV_{safe_co}_{safe_t}_{ts}.pdf"
        filepath = str(self.OUTPUT_DIR / filename)

        doc = SimpleDocTemplate(filepath, pagesize=A4,
                                leftMargin=18*mm, rightMargin=18*mm,
                                topMargin=14*mm, bottomMargin=14*mm)

        personal = profile.get("personal", {})
        summaries = profile.get("summaries", {})
        summary = summaries.get(summary_type, summaries.get("ai_product_manager", ""))
        experience = profile.get("experience", [])
        projects = profile.get("projects", [])
        education = profile.get("education", [])
        skills = profile.get("skills", {})
        certs = profile.get("certificates", [])
        languages = profile.get("languages", [])
        references = profile.get("references", [])

        # Pick best summary variant based on job title
        jt = job_title.lower()
        if any(k in jt for k in ["account", "sales", "growth", "revenue"]):
            summary = summaries.get("account_manager", summary)
        elif any(k in jt for k in ["delivery", "program"]):
            summary = summaries.get("delivery_manager", summary)
        elif any(k in jt for k in ["finance", "analyst", "accounting"]):
            summary = summaries.get("finance_analyst", summary)

        story = []
        s = self.styles

        # ── HEADER ──
        story.append(Paragraph(personal.get("name", "").upper(), s["name"]))
        story.append(Paragraph(f"{job_title}", s["title"]))
        contact_parts = [
            f"📞 {personal.get('phone', '')}",
            f"📧 {personal.get('email', '')}",
            f"📍 {personal.get('location', '')}",
            f"🔗 {personal.get('linkedin', '')}",
        ]
        story.append(Paragraph("  |  ".join(contact_parts), s["contact"]))
        story.append(HRFlowable(width="100%", thickness=0.8, color=HexColor("#1a1a3e")))

        # ── PROFESSIONAL SUMMARY ──
        story.append(Paragraph("PROFESSIONAL SUMMARY", s["section"]))
        story.append(Paragraph(summary.strip(), s["body"]))
        story.append(Spacer(1, 3))

        # ── SKILLS & TOOLS ──
        story.append(Paragraph("SKILLS & TOOLS", s["section"]))
        # Select relevant skill categories based on job
        skill_order = ["delivery_execution", "ai_automation", "data_platforms", "leadership"]
        if any(k in jt for k in ["finance", "analyst", "accounting"]):
            skill_order = ["finance", "delivery_execution", "data_platforms", "technical"]
        for cat in skill_order:
            items = skills.get(cat, [])
            if items:
                label = cat.replace("_", " ").title()
                story.append(Paragraph(
                    f"<b>{label}:</b> {', '.join(items)}", s["skills"]
                ))
        story.append(Spacer(1, 2))

        # ── PROJECTS / PRODUCTS ──
        story.append(Paragraph("PROJECTS / PRODUCTS", s["section"]))
        for proj in projects:
            block = []
            block.append(Paragraph(
                f"<b>{proj['name']}</b> ({proj.get('role', '')})", s["subsection"]
            ))
            block.append(Paragraph(
                f"{proj.get('company', '')}  |  {proj.get('period', '')}", s["meta"]
            ))
            for bullet in proj.get("bullets", []):
                block.append(Paragraph(f"• {bullet}", s["bullet"]))
            block.append(Spacer(1, 3))
            story.append(KeepTogether(block))

        # ── PROFESSIONAL EXPERIENCE ──
        story.append(Paragraph("PROFESSIONAL EXPERIENCE", s["section"]))
        for exp in experience:
            block = []
            title_text = exp.get("alt_title", exp["title"])
            block.append(Paragraph(f"<b>{title_text}</b>", s["subsection"]))
            block.append(Paragraph(
                f"{exp['company']}  |  {exp.get('location', '')}  |  {exp['period']}", s["meta"]
            ))
            for bullet in exp.get("bullets", [])[:6]:  # Top 6 bullets per role
                block.append(Paragraph(f"• {bullet}", s["bullet"]))
            block.append(Spacer(1, 3))
            story.append(KeepTogether(block))

        # ── EDUCATION ──
        story.append(Paragraph("EDUCATION", s["section"]))
        for edu in education:
            text = f"<b>{edu['degree']}</b>  —  {edu['institution']}  |  {edu['period']}"
            if edu.get("note"):
                text += f"  <i>({edu['note']})</i>"
            story.append(Paragraph(text, s["body"]))
        story.append(Spacer(1, 3))

        # ── CERTIFICATIONS ──
        story.append(Paragraph("CERTIFICATIONS", s["section"]))
        cert_text = "  •  ".join(certs)
        story.append(Paragraph(cert_text, s["body"]))
        story.append(Spacer(1, 3))

        # ── LANGUAGES ──
        story.append(Paragraph("LANGUAGES", s["section"]))
        lang_parts = [f"{l['lang']} ({l['level']})" for l in languages]
        story.append(Paragraph("  |  ".join(lang_parts), s["body"]))

        # ── REFERENCES ──
        if references:
            story.append(Spacer(1, 4))
            story.append(Paragraph("REFERENCES", s["section"]))
            for ref in references:
                story.append(Paragraph(
                    f"<b>{ref['name']}</b> — {ref['title']}, {ref['company']}  |  {ref['email']}",
                    s["body"]
                ))

        doc.build(story)
        return filepath

    def generate_docx(self, profile: dict, job_title: str, company: str,
                      summary_type: str = "ai_product_manager") -> str:
        """Generate a detailed DOCX CV."""
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        safe_co = "".join(c if c.isalnum() else "_" for c in company)[:25]
        safe_t = "".join(c if c.isalnum() else "_" for c in job_title)[:25]
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"CV_{safe_co}_{safe_t}_{ts}.docx"
        filepath = str(self.OUTPUT_DIR / filename)

        doc = Document()
        style = doc.styles["Normal"]
        style.font.size = Pt(10)
        style.font.name = "Calibri"

        personal = profile.get("personal", {})
        summaries = profile.get("summaries", {})
        summary = summaries.get(summary_type, "")

        # Name
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(personal.get("name", "").upper())
        run.bold = True
        run.font.size = Pt(20)

        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(job_title)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Contact
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact = f"{personal.get('phone','')} | {personal.get('email','')} | {personal.get('location','')} | {personal.get('linkedin','')}"
        run = p.add_run(contact)
        run.font.size = Pt(9)

        # Summary
        doc.add_heading("PROFESSIONAL SUMMARY", level=2)
        doc.add_paragraph(summary.strip())

        # Skills
        doc.add_heading("SKILLS & TOOLS", level=2)
        for cat, items in profile.get("skills", {}).items():
            label = cat.replace("_", " ").title()
            p = doc.add_paragraph()
            run = p.add_run(f"{label}: ")
            run.bold = True
            p.add_run(", ".join(items))

        # Projects
        doc.add_heading("PROJECTS / PRODUCTS", level=2)
        for proj in profile.get("projects", []):
            p = doc.add_paragraph()
            run = p.add_run(f"{proj['name']} ({proj.get('role', '')})")
            run.bold = True
            p = doc.add_paragraph()
            run = p.add_run(f"{proj.get('company', '')} | {proj.get('period', '')}")
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run.font.size = Pt(9)
            for bullet in proj.get("bullets", []):
                doc.add_paragraph(bullet, style="List Bullet")

        # Experience
        doc.add_heading("PROFESSIONAL EXPERIENCE", level=2)
        for exp in profile.get("experience", []):
            p = doc.add_paragraph()
            run = p.add_run(exp.get("alt_title", exp["title"]))
            run.bold = True
            p = doc.add_paragraph()
            run = p.add_run(f"{exp['company']} | {exp.get('location','')} | {exp['period']}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            for bullet in exp.get("bullets", [])[:6]:
                doc.add_paragraph(bullet, style="List Bullet")

        # Education
        doc.add_heading("EDUCATION", level=2)
        for edu in profile.get("education", []):
            text = f"{edu['degree']} — {edu['institution']} | {edu['period']}"
            doc.add_paragraph(text)

        # Certs
        doc.add_heading("CERTIFICATIONS", level=2)
        for cert in profile.get("certificates", []):
            doc.add_paragraph(cert, style="List Bullet")

        # Languages
        doc.add_heading("LANGUAGES", level=2)
        langs = [f"{l['lang']} ({l['level']})" for l in profile.get("languages", [])]
        doc.add_paragraph(" | ".join(langs))

        doc.save(filepath)
        return filepath
