"""CV Generator — Creates ATS-optimized PDF and DOCX CVs tailored per job."""
import json
import os
from datetime import datetime
from pathlib import Path
from io import BytesIO

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, mm
from reportlab.lib.colors import HexColor
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from cv_parser import call_deepseek

OUTPUT_DIR = Path(__file__).parent / "data" / "tailored_cvs"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


async def generate_tailored_cv(user_profile: dict, job: dict,
                               provider: str = "deepseek", api_key: str = "") -> dict:
    """Generate best-in-class ATS-optimized CV tailored for a specific job."""

    prompt = f"""You are a world-class resume writer who has helped thousands of candidates land jobs at 
top companies. Create an ATS-optimized resume that will score 95%+ on any ATS system.

TARGET JOB: {job.get('title', '')} at {job.get('company', '')}

JOB DESCRIPTION:
{(job.get('description', '') or '')[:4000]}

CANDIDATE DATA:
Name: {user_profile.get('name', '')}
Email: {user_profile.get('email', '')}
Phone: {user_profile.get('phone', '')}
Location: {user_profile.get('location', '')}
LinkedIn: {user_profile.get('linkedin', '')}

Current Summary: {user_profile.get('summary', '')}

Skills: {json.dumps(user_profile.get('skills', {}))}

Experience:
{json.dumps(user_profile.get('experience', []), indent=2)}

Education:
{json.dumps(user_profile.get('education', []), indent=2)}

Certifications: {json.dumps(user_profile.get('certifications', []))}

Projects:
{json.dumps(user_profile.get('projects', []), indent=2)}

STRICT RULES FOR BEST-IN-CLASS ATS RESUME:

FORMAT:
- Single column layout ONLY (no tables, no columns, no graphics)
- Reverse chronological order
- Standard section headers in CAPS: PROFESSIONAL SUMMARY, EXPERIENCE, SKILLS, EDUCATION, CERTIFICATIONS, PROJECTS
- Use simple bullet points with "•" character
- Keep to exactly 2 pages

CONTENT OPTIMIZATION:
1. TITLE LINE: Make it match or closely mirror the job title (e.g., if job is "Senior Product Manager - AI", use that as the title under the name)
2. PROFESSIONAL SUMMARY: 3-4 lines that mirror the job description language. Include years of experience, key domain expertise, and 2-3 quantified achievements. Use exact keywords from the job posting.
3. EXPERIENCE: 
   - Each role: Title | Company | Location | Date Range
   - 4-6 bullet points per role, starting with strong action verbs
   - EVERY bullet must have a quantified metric (%, $, #, time saved)
   - Front-load bullets most relevant to the TARGET job
   - Weave in keywords from the job description naturally
4. SKILLS: Group into Technical Skills, Tools & Platforms, Methodologies, Soft Skills
   - Include EXACT skill names from the job description
   - For missing skills, include adjacent/related skills the candidate has
5. EDUCATION: Degree | Institution | Year
6. CERTIFICATIONS: List all, most relevant first
7. PROJECTS: Only include if highly relevant to the target job

KEYWORD STRATEGY:
- Extract the top 15 keywords from the job description
- Ensure each keyword appears at least once in the resume
- Place keywords in Summary, Experience bullets, AND Skills section
- Use both the full term and acronym (e.g., "Artificial Intelligence (AI)")

Return ONLY the complete resume text. Start with the candidate's name on line 1, contact info on line 2, then sections.
Do NOT include any markdown formatting, just plain text.
"""

    cv_text = await call_deepseek(prompt, system="You are a world-class resume writer. Return only the resume text, no explanations.",
                                provider=provider, api_key=api_key)
    if not cv_text or len(cv_text) < 100:
        return {"error": "Failed to generate CV"}

    safe_company = "".join(c if c.isalnum() else "_" for c in job.get("company", "")[:30])
    safe_title = "".join(c if c.isalnum() else "_" for c in job.get("title", "")[:30])
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = f"CV_{safe_company}_{safe_title}_{ts}"

    pdf_path = str(OUTPUT_DIR / f"{base_name}.pdf")
    docx_path = str(OUTPUT_DIR / f"{base_name}.docx")

    _generate_pdf(cv_text, pdf_path, user_profile)
    _generate_docx(cv_text, docx_path, user_profile)

    return {"pdf_path": pdf_path, "docx_path": docx_path, "cv_text": cv_text}


def _generate_pdf(cv_text: str, output_path: str, profile: dict):
    """Generate a clean ATS-friendly PDF from CV text."""
    doc = SimpleDocTemplate(output_path, pagesize=A4,
                            leftMargin=0.6*inch, rightMargin=0.6*inch,
                            topMargin=0.5*inch, bottomMargin=0.5*inch)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CVName", fontSize=16, fontName="Helvetica-Bold",
                              alignment=TA_CENTER, spaceAfter=4))
    styles.add(ParagraphStyle(name="CVContact", fontSize=9, fontName="Helvetica",
                              alignment=TA_CENTER, spaceAfter=8, textColor=HexColor("#444444")))
    styles.add(ParagraphStyle(name="CVSection", fontSize=11, fontName="Helvetica-Bold",
                              spaceAfter=4, spaceBefore=10, textColor=HexColor("#1a1a2e")))
    styles.add(ParagraphStyle(name="CVBody", fontSize=9.5, fontName="Helvetica",
                              spaceAfter=3, leading=13))
    styles.add(ParagraphStyle(name="CVBullet", fontSize=9.5, fontName="Helvetica",
                              spaceAfter=2, leading=12, leftIndent=15))

    story = []
    lines = cv_text.strip().split("\n")
    section_headers = {"SUMMARY", "EXPERIENCE", "EDUCATION", "SKILLS", "CERTIFICATIONS",
                       "PROJECTS", "LANGUAGES", "CERTIFICATES", "PROFESSIONAL EXPERIENCE",
                       "TECHNICAL SKILLS", "KEY SKILLS", "WORK EXPERIENCE"}

    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            story.append(Spacer(1, 4))
            continue
        if i == 0:
            story.append(Paragraph(line, styles["CVName"]))
            continue
        if i <= 2 and ("@" in line or "+" in line or "linkedin" in line.lower()):
            story.append(Paragraph(line, styles["CVContact"]))
            continue

        upper = line.upper().rstrip(":")
        if upper in section_headers:
            story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#cccccc")))
            story.append(Paragraph(upper, styles["CVSection"]))
            continue

        if line.startswith("- ") or line.startswith("• "):
            story.append(Paragraph(f"• {line[2:]}", styles["CVBullet"]))
        else:
            story.append(Paragraph(line, styles["CVBody"]))

    doc.build(story)


def _generate_docx(cv_text: str, output_path: str, profile: dict):
    """Generate a clean ATS-friendly DOCX from CV text."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    section_headers = {"SUMMARY", "EXPERIENCE", "EDUCATION", "SKILLS", "CERTIFICATIONS",
                       "PROJECTS", "LANGUAGES", "CERTIFICATES", "PROFESSIONAL EXPERIENCE",
                       "TECHNICAL SKILLS", "KEY SKILLS", "WORK EXPERIENCE"}

    lines = cv_text.strip().split("\n")
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue
        if i == 0:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(16)
            continue
        if i <= 2 and ("@" in line or "+" in line or "linkedin" in line.lower()):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)
            continue

        upper = line.upper().rstrip(":")
        if upper in section_headers:
            p = doc.add_paragraph()
            run = p.add_run(upper)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(26, 26, 46)
            continue

        if line.startswith("- ") or line.startswith("• "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    doc.save(output_path)


COVER_LETTER_DIR = Path(__file__).parent / "data" / "cover_letters"
COVER_LETTER_DIR.mkdir(parents=True, exist_ok=True)


async def generate_cover_letter(user_profile: dict, job: dict,
                                 provider: str = "deepseek", api_key: str = "") -> dict:
    """Generate a personalized cover letter tailored for a specific job."""

    prompt = f"""You are a professional career coach. Write a compelling, personalized cover letter.

TARGET JOB: {job.get('title', '')} at {job.get('company', '')}
LOCATION: {job.get('location', '')}

JOB DESCRIPTION:
{(job.get('description', '') or '')[:3000]}

CANDIDATE:
Name: {user_profile.get('name', '')}
Email: {user_profile.get('email', '')}
Phone: {user_profile.get('phone', '')}
Location: {user_profile.get('location', '')}
LinkedIn: {user_profile.get('linkedin', '')}

Summary: {user_profile.get('summary', '')}
Skills: {json.dumps(user_profile.get('skills', {}))}
Experience: {json.dumps(user_profile.get('experience', [])[:3], indent=2)}

RULES:
1. Address it to "Hiring Manager" at {job.get('company', '')}
2. Opening paragraph: Express enthusiasm for the specific role and company. Mention how you found the role.
3. Body paragraph 1: Highlight 2-3 most relevant experiences that directly match the job requirements. Use specific metrics and achievements.
4. Body paragraph 2: Connect your skills to the company's needs. Show you've researched the company.
5. Closing paragraph: Express eagerness to discuss further, mention availability.
6. Keep it to 1 page (300-400 words max).
7. Professional but warm tone.
8. Do NOT use generic phrases like "I am writing to apply for..."
9. Start with the candidate's contact info, then date, then the letter.

Return ONLY the cover letter text, no markdown formatting.
"""

    cl_text = await call_deepseek(prompt,
        system="You are a professional career coach. Return only the cover letter text.",
        provider=provider, api_key=api_key)

    if not cl_text or len(cl_text) < 100:
        return {"error": "Failed to generate cover letter"}

    safe_company = "".join(c if c.isalnum() else "_" for c in job.get("company", "")[:30])
    safe_title = "".join(c if c.isalnum() else "_" for c in job.get("title", "")[:30])
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = f"CL_{safe_company}_{safe_title}_{ts}"

    pdf_path = str(COVER_LETTER_DIR / f"{base_name}.pdf")
    docx_path = str(COVER_LETTER_DIR / f"{base_name}.docx")

    _generate_cover_letter_pdf(cl_text, pdf_path, user_profile)
    _generate_cover_letter_docx(cl_text, docx_path, user_profile)

    return {"pdf_path": pdf_path, "docx_path": docx_path, "text": cl_text}


def _generate_cover_letter_pdf(text: str, output_path: str, profile: dict):
    """Generate a clean cover letter PDF."""
    doc = SimpleDocTemplate(output_path, pagesize=A4,
                            leftMargin=0.8*inch, rightMargin=0.8*inch,
                            topMargin=0.7*inch, bottomMargin=0.7*inch)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CLBody", fontSize=11, fontName="Helvetica",
                              spaceAfter=12, leading=16))
    styles.add(ParagraphStyle(name="CLName", fontSize=13, fontName="Helvetica-Bold",
                              spaceAfter=4))
    styles.add(ParagraphStyle(name="CLContact", fontSize=9, fontName="Helvetica",
                              spaceAfter=2, textColor=HexColor("#444444")))

    story = []
    lines = text.strip().split("\n")
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            story.append(Spacer(1, 8))
            continue
        if i == 0:
            story.append(Paragraph(line, styles["CLName"]))
        elif i <= 3 and ("@" in line or "+" in line or "linkedin" in line.lower()):
            story.append(Paragraph(line, styles["CLContact"]))
        else:
            story.append(Paragraph(line, styles["CLBody"]))

    doc.build(story)


def _generate_cover_letter_docx(text: str, output_path: str, profile: dict):
    """Generate a clean cover letter DOCX."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = text.strip().split("\n")
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue
        if i == 0:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(13)
        elif i <= 3 and ("@" in line or "+" in line or "linkedin" in line.lower()):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)
        else:
            doc.add_paragraph(line)

    doc.save(output_path)
